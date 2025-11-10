# generator.py
import os
import chromadb
from chromadb.utils import embedding_functions
import openai
from docx import Document
import io
from datetime import datetime

# Initialize OpenAI
from dotenv import load_dotenv
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# Initialize ChromaDB
client = chromadb.PersistentClient(path="./chroma_db")
collection = client.get_or_create_collection(name="marketing_templates")

# Embedding function
embedding_function = embedding_functions.SentenceTransformerEmbeddingFunction(
    model_name="all-MiniLM-L6-v2"
)

def extract_text_from_docx(filepath):
    try:
        doc = Document(filepath)
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
        return ""  # Or handle the error as needed

# Load template documents into Chroma (only once)
def load_templates():
    if collection.count() > 0:
        return

    brand_template = extract_text_from_docx("Brand Guideline.docx")
    digital_template = extract_text_from_docx("Digital Strategy.docx")
    digital_example = extract_text_from_docx("Digital Strategy Example.docx")

    texts = [brand_template, digital_template, digital_example]
    metadatas = [
        {"type": "brand_guideline_template"},
        {"type": "digital_strategy_template"},
        {"type": "digital_strategy_example"}
    ]
    ids = ["brand_template", "digital_template", "digital_example"]

    collection.add(
        documents=texts,
        metadatas=metadatas,
        ids=ids,
        embeddings=embedding_function(texts)
    )

# RAG Retrieval
def retrieve_context(query, agent_type):
    # results = collection.query(
    #     query_texts=[query],
    #     n_results=3,
    #     where={"type": {"$in": [agent_type, "digital_strategy_example"] if "digital" in agent_type else [agent_type]}}
    # )
    # return "\n\n".join(results['documents'][0])

    where_clause = {"type": agent_type}
    if agent_type == "digital_strategy_template":
        where_clause = {"type": {"$in": ["digital_strategy_template", "digital_strategy_example"]}}
    if agent_type == "brand_strategy_template":
        where_clause = {"type": {"$in": ["brand_strategy_template"]}}

    results = collection.query(
        query_texts=[query],
        n_results=3,
        where=where_clause
    )
    docs = results['documents'][0]
    return "\n\n".join(docs) if docs else ""

# Supercharged Prompt Templates
PROMPT_TEMPLATES = {
    "brand": """
You are a senior brand strategist at LFTFIELD, a premium marketing agency. 
Your task is to generate a **Brand Strategy & Guideline Document** based on the client inputs.

Use the retrieved context as reference for structure, tone, and best practices.
Follow this exact structure:

1. **Brand Overview**
   - Mission, Vision, Values
   - Brand Story

2. **Target Audience**
   - Primary & Secondary Personas (demographics, psychographics)

3. **Brand Positioning**
   - Unique Value Proposition
   - Positioning Statement

4. **Brand Voice & Tone**
   - Voice adjectives
   - Do's and Don'ts

5. **Visual Identity**
   - Logo usage
   - Color palette
   - Typography
   - Imagery style

6. **Brand Applications**
   - Examples across touchpoints

Use professional, confident, and concise language. 
Incorporate insights from retrieved examples where relevant.

Client Input:
{input}

Retrieved Context:
{context}

Output only the final document in clean Markdown.
""",

    "digital": """
You are a senior digital strategist at LFTFIELD. 
Generate a **Comprehensive Digital Marketing Strategy** using the Flywheel Framework (Attract → Engage → Delight).

Use SMART goals. Include:
- Company Background
- Products/Services
- Marketing Goals (SMART)
- SWOT
- Competitive Analysis
- Target Customers + Buying Cycle
- USP + Brand Relevance
- Flywheel Strategy (Attract/Engage/Delight + Friction)
- Preliminary Projects
- KPIs per channel
- Channel Status & Requirements
- Budget Recommendations
- Next Steps

Use the MMTTC example as a structural and quality benchmark.
Be concise, data-driven, and actionable.

Client Input:
{input}

Retrieved Context:
{context}

Output in clean Markdown with clear headings and bullets.
"""
}

# Generate Draft
def generate_draft(agent_type, answers):
    load_templates()
    
    # Format answers
    input_summary = "\n".join([f"Q: {q}\nA: {a}" for q, a in answers.items()])
    
    # Retrieve relevant context
    context = retrieve_context(input_summary, agent_type)
    
    # Final prompt
    prompt = PROMPT_TEMPLATES[agent_type].format(input=input_summary, context=context)
    
    response = openai.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "system", "content": prompt}],
        temperature=0.3,
        max_tokens=4000
    )
    
    return response.choices[0].message.content

# Export to Word
def export_to_word(markdown_text, client_name):
    doc = Document()
    doc.add_heading(f"{client_name} - Strategy Document", 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph()

    for line in markdown_text.split("\n"):
        line = line.strip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("· "):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line[2:])
        elif line.startswith("1. "):
            p = doc.add_paragraph(style='List Number')
            p.add_run(line[3:])
        elif line:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
