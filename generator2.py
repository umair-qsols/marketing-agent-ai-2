import os
import chromadb
from chromadb.utils import embedding_functions
import openai
from docx import Document
import io
from datetime import datetime

# Initialize OpenAI
from dotenv import load_dotenv
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# Initialize ChromaDB
client = chromadb.PersistentClient(path="./chroma_db")
collection = client.get_or_create_collection(name="marketing_templates")

# Embedding function
embedding_function = embedding_functions.SentenceTransformerEmbeddingFunction(
    model_name="all-MiniLM-L6-v2"
)

def extract_text_from_docx(filepath):
    try:
        doc = Document(filepath)
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
        return ""

# Load template documents into Chroma (only once)
def load_templates():
    if collection.count() > 0:
        print(f"Templates already loaded. Count: {collection.count()}")
        return

    print("Loading templates into ChromaDB...")
    
    brand_template = extract_text_from_docx("Brand Guideline.docx")
    digital_template = extract_text_from_docx("Digital Strategy.docx")
    digital_example = extract_text_from_docx("Digital Strategy Example.docx")

    # Verify content was extracted
    if not brand_template:
        print("WARNING: Brand template is empty!")
    if not digital_template:
        print("WARNING: Digital template is empty!")
    if not digital_example:
        print("WARNING: Digital example is empty!")

    texts = [brand_template, digital_template, digital_example]
    metadatas = [
        {"type": "brand_template"},
        {"type": "digital_template"},
        {"type": "digital_example"}
    ]
    ids = ["brand_template", "digital_template", "digital_example"]

    collection.add(
        documents=texts,
        metadatas=metadatas,
        ids=ids,
        embeddings=embedding_function(texts)
    )
    
    print(f"Templates loaded successfully. Total count: {collection.count()}")

# RAG Retrieval
def retrieve_context(query, agent_type):
    """
    agent_type should be either 'brand' or 'digital'
    """
    print(f"Retrieving context for agent_type: {agent_type}")
    
    # Map agent type to template types
    if agent_type == "digital":
        where_clause = {"type": {"$in": ["digital_template", "digital_example"]}}
    elif agent_type == "brand":
        where_clause = {"type": "brand_template"}
    else:
        print(f"Unknown agent_type: {agent_type}")
        return ""

    try:
        results = collection.query(
            query_texts=[query],
            n_results=3,
            where=where_clause
        )
        
        docs = results['documents'][0] if results['documents'] else []
        print(f"Retrieved {len(docs)} documents")
        
        if docs:
            context = "\n\n".join(docs)
            print(f"Context length: {len(context)} characters")
            return context
        else:
            print("No documents retrieved!")
            return ""
            
    except Exception as e:
        print(f"Error during retrieval: {e}")
        return ""

# Supercharged Prompt Templates
PROMPT_TEMPLATES = {
    "brand": """
You are a senior brand strategist at LFTFIELD, a premium marketing agency. 
Your task is to generate a **Brand Strategy & Guideline Document** based on the client inputs.

**CRITICAL INSTRUCTION**: You MUST follow the EXACT structure from the template below. Do NOT create a generic brand guideline. Use this EXACT outline:

REQUIRED SECTIONS (DO NOT SKIP ANY):

# Company Description
[Detailed overview of the company, its background, industry position, and what they do]

# Brand Wheel
One essential part of the brand development process is the "brand wheel," a templated approach to understanding your brand by breaking it down into five categories:

## Attributes
[List 3-5 key brand attributes]

## Benefits
[List 3-5 benefits the brand provides to customers]

## Values
[List 3-5 core brand values]

## Personality
[List 3-5 personality traits of the brand]

## Essence
[Single phrase or word that captures the brand essence]

# Audience Personas
Your Audience Personas should epitomize your customer base. These fictional profiles will help to ensure your brand and marketing efforts will appeal to your audience.

[Create 2-3 detailed personas with:]
- Demographics (age, education, location, income)
- Background
- Goals
- Challenges
- Motivations

# Competitor Research
Profiling your competitors gives people a unique insight into your industry. 

[Analyze 2-3 key competitors:]
- Company name and overview
- What they offer
- Their strengths
- Key learnings

# Brand Positioning
Brand positioning is the process of placing your brand in the minds of your customers.

**Using the following formula:** [Brand Name]'s [offering] is the only [category/service/product] that [benefit you bring to your customers].

[Write the positioning statement]

# Brand Story
Your Brand Story is unique to you -- it can be funny, unexpected, serious, ambitious... but one thing is for sure: it must spark an emotional reaction. Include the direct pain points that you solve for customers.

[Write a compelling 2-3 paragraph brand story]

# Brand Values
You will find your Brand Values at the core of your Brand Strategy. State each value with a descriptive sentence. Avoid clichés like "transparent" and "honest."

[List 3-5 values with descriptions]

# Brand Mission
Detail exactly where your brand is going and what you want to achieve.

[Write mission statement as a paragraph]

# Brand Touchpoints
A Brand Touch Point is the time and place where a customer comes in contact with your brand.

[List and describe 5-8 touchpoints where customers interact with the brand]

# Brand Messaging
Your Brand Messaging is "what" you're trying to communicate and how you communicate it.

[List 3-5 key messages]
[Optional: Break down into brand pillars]

# Tone of Voice
Your Tone of Voice describes how your brand communicates with the audience and thus influences how people perceive your messaging.

Create a table with:
| Characteristic | Description | Do's | Don'ts |
[Fill in 3-5 rows]

---

Retrieved Template Context for Reference:
{context}

---

Client Input:
{input}

---

**FINAL INSTRUCTIONS:**
1. You MUST include ALL sections listed above - no exceptions
2. Follow the exact heading structure (# for main sections, ## for subsections)
3. Provide substantial content for each section - not just placeholders
4. Use professional, strategic language
5. Base all content on the client input provided
6. Output ONLY the completed document in Markdown format
7. Do NOT add any preamble or explanation - start directly with "# Company Description"

BEGIN OUTPUT NOW:
""",

    "digital": """
You are a senior digital strategist at LFTFIELD. 
Generate a **Comprehensive Digital Marketing Strategy** using the Flywheel Framework (Attract → Engage → Delight).

**IMPORTANT**: Use the retrieved context below (especially the MMTTC example) as your PRIMARY reference for structure, depth, and quality.

Retrieved Template Context:
{context}

---

Now, using the structure and quality standards from the examples above, create a complete digital strategy for this client:

Client Input:
{input}

**Required Structure:**
1. INTRODUCTION
   - Company Background
   - Products & Services
   - Marketing Goals (SMART format)

2. RESEARCH & ANALYSIS
   - SWOT Analysis
   - Competitive Analysis
   - Target Customers (detailed personas)
   - Buying Cycle
   - Unique Selling Proposition
   - Brand Relevance

3. OUR STRATEGY
   - Proposed Strategy (Flywheel: Attract/Engage/Delight)
   - Preliminary Projects
   - Key Performance Indicators (by channel)
   - Marketing Channels

4. THE ROAD AHEAD
   - Summary
   - Next Steps

Use professional, data-driven language. Be specific and actionable.
Output in clean Markdown with clear headings and bullet points.
"""
}

# Generate Draft
def generate_draft(agent_type, answers):
    """
    agent_type: 'brand' or 'digital'
    answers: dict of {question: answer}
    """
    print(f"\n=== Generating {agent_type} draft ===")
    
    # Ensure templates are loaded
    load_templates()
    
    # Format answers
    input_summary = "\n".join([f"Q: {q}\nA: {a}" for q, a in answers.items()])
    print(f"Input summary length: {len(input_summary)} characters")
    
    # Retrieve relevant context
    context = retrieve_context(input_summary, agent_type)
    
    if not context:
        print("WARNING: No context retrieved! Output quality may be poor.")
    
    # Get the appropriate prompt template
    if agent_type not in PROMPT_TEMPLATES:
        raise ValueError(f"Unknown agent_type: {agent_type}. Must be 'brand' or 'digital'")
    
    # Final prompt
    prompt = PROMPT_TEMPLATES[agent_type].format(input=input_summary, context=context)
    print(f"Final prompt length: {len(prompt)} characters")
    
    print("Calling OpenAI API...")
    # Use different token limits based on agent type
    max_tokens = 6000 if agent_type == "brand" else 4000
    
    response = openai.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "system", "content": prompt}],
        temperature=0.3,
        max_tokens=max_tokens
    )
    
    generated_text = response.choices[0].message.content
    print(f"Generated text length: {len(generated_text)} characters")
    print("=== Generation complete ===\n")
    
    return generated_text

# Export to Word
def export_to_word(markdown_text, client_name):
    doc = Document()
    doc.add_heading(f"{client_name} - Strategy Document", 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph()

    for line in markdown_text.split("\n"):
        line = line.strip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("· "):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line[2:])
        elif line.startswith("1. "):
            p = doc.add_paragraph(style='List Number')
            p.add_run(line[3:])
        elif line:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Utility function to reset/debug ChromaDB
def reset_chroma():
    """Call this if you need to force reload templates"""
    global collection
    client.delete_collection(name="marketing_templates")
    collection = client.get_or_create_collection(name="marketing_templates")
    load_templates()
    print("ChromaDB reset and templates reloaded")